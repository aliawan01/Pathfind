import docx
import os

from docx.shared import Inches, Cm

from color_manager import *
from ui_manager import *
from pathfinding_algorithms import *
from maze_generation_algorithms import *
from animations import *
from main import *
from grid import *
from queue_classes import *
from stack import *
from networking import *
import re

import inspect
from types import FunctionType
import itertools
import sys


# is_present = False
# document = docx.Document()
#
# try:
#     document = docx.Document("..\\generated_table.docx")
#     is_present =True
# except:
#     pass
#
# if is_present:
#     os.remove('..\\generated_table.docx')
#     document = docx.Document()

document = docx.Document('..\\generated_table.docx')
sections = document.sections

# document.add_heading("Level 3 heading", level=3)
#
# table = document.add_table(rows=5, cols=4, style='Table Grid')
# row_cells = table.rows[0].cells
# row_cells[0].text = 'Name'
# row_cells[1].text = 'Parameters'
# row_cells[2].text = 'Return Value'
# row_cells[3].text = 'Description'

def listMethods(cls):
    return set(x for x, y in cls.__dict__.items()
               if isinstance(y, (FunctionType, classmethod, staticmethod)))

def listParentMethods(cls):
    return set(itertools.chain.from_iterable(
        listMethods(c).union(listParentMethods(c)) for c in cls.__bases__))

def list_subclass_methods(cls,is_narrow):
    methods = listMethods(cls)
    if  is_narrow:
        parentMethods = listParentMethods(cls)
        return set(cls for cls in methods if not (cls in parentMethods))
    else:
        return methods

def get_obj_of_function_name(class_obj, function_name):
    for name, obj in inspect.getmembers(class_obj, inspect.isfunction):
        if name == function_name:
            return obj

def parse_and_print_function_obj(function_obj):
    documentation = function_obj.__doc__
    if documentation != None:
        # documentation = documentation.replace('\n', '')
        documentation = documentation.replace('  ', '')

        param_matches = []
        return_matches = []
        for line in documentation.split('\n'):
            if line.find("@param") == 0:
                param_matches.append(line)
            elif line.find("@return") == 0:
                return_matches.append(line)

        param_string = ""
        for param in param_matches:
            param = param.replace("@param ", '')
            param += ", "
            param_string += param

        if param_string != "":
            param_string = param_string[:-2]

        if param_string == "":
            param_string = "None"

        return_string = ""
        if return_matches != []:
            return_string = return_matches[0].replace("@return:", "")
            if return_string[0] == ' ':
                return_string = return_string[1:]

        if return_string == "":
            return_string = "None"

        for p_match in param_matches:
            documentation = documentation.replace(p_match, "")

        for r_match in return_matches:
            documentation = documentation.replace(r_match, "")

        documentation = documentation.replace("\n", " ")
        documentation = documentation.strip()
        if documentation[-1] != '.':
            documentation += '.'
        # if documentation[0] == ' ':
        #     documentation = documentation[1:]


        return {'documentation': documentation, 'param': param_string, 'return': return_string}
    else:
        return {'documentation': '', 'param': '', 'return': ''}

class_names_and_objs = inspect.getmembers(sys.modules[__name__], inspect.isclass)

table = document.add_table(rows=1, cols=4, style='pathfinding_style')
table.autofit = False
original_row = table.rows[0].cells
original_row[0].text = 'Name'
original_row[1].text = 'Parameters'
original_row[2].text = 'Return Value'
original_row[3].text = 'Description'

new_row = table.add_row().cells
new_row[0].text = 'main'
new_row[1].text = 'None'
new_row[2].text = 'None'
new_row[3].text = 'This is the main function which will run the game.'

for cell in table.columns[0].cells:
    cell.width = Cm(4)

for cell in table.columns[1].cells:
    cell.width = Cm(4)

for cell in table.columns[2].cells:
    cell.width = Cm(3.3)

for cell in table.columns[3].cells:
    cell.width = Cm(6.7)

# TODO(ali): Find a way to handle the IntEnums seperately and also make a different table for them as well.
for class_name, class_obj in class_names_and_objs:
    if class_name in ['IntEnum', 'Enum', 'Cm', 'Inches', 'Color', 'FunctionType']:
        continue

    if IntEnum in class_obj.__bases__:
        continue
        # print(class_name)
        # for name in class_obj._member_names_:
        #     print("           ", name)
        #     exec("print(class_obj." + name + ")")

    document.add_paragraph("")
    document.add_heading(class_name + ' Class', 3)

    table = document.add_table(rows=1, cols=4, style='pathfinding_style')
    table.autofit = False
    # print(table.columns[3].width/360000)
    # table.columns[0].width = Cm(1)
    # table.columns[1].width = Cm(1)
    # table.columns[2].width = Cm(1)
    # table.columns[3].width = Cm(7)
    row = table.rows[0].cells
    row[0].text = 'Name'
    row[1].text = 'Parameters'
    row[2].text = 'Return Value'
    row[3].text = 'Description'

    class_info = []
    subclasses_list = list_subclass_methods(class_obj, False)

    for function_name in subclasses_list:
        function_obj = get_obj_of_function_name(class_obj, function_name)
        doc_dict = parse_and_print_function_obj(function_obj)
        class_info.append((function_name, doc_dict['param'], doc_dict['return'], doc_dict['documentation']))

    # print(class_info)
    for name, param, return_val, description in class_info:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = param
        row[2].text = return_val
        row[3].text = description

    for cell in table.columns[0].cells:
        cell.width = Cm(4)

    for cell in table.columns[1].cells:
        cell.width = Cm(4)

    for cell in table.columns[2].cells:
        cell.width = Cm(3.3)

    for cell in table.columns[3].cells:
        cell.width = Cm(6.7)


document.save('..\\generated_table.docx')